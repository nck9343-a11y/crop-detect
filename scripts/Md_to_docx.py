"""
把 markdown 草稿转成 Word 文档
================================
只处理本项目草稿实际用到的语法：标题、段落、表格、无序列表、引用块，
以及行内的 **加粗** 与 `代码`。刻意不做成通用转换器——
通用实现要处理的边界情况远多于此，而这里只需要稳定处理一份已知格式的文件。

用法：
    python scripts/Md_to_docx.py            # 默认 paper/论文.md -> paper/论文.docx
    python scripts/Md_to_docx.py 输入.md 输出.docx
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# 按语言选字体：英文正文若沿用宋体，一眼即可看出是从中文模板生成的。
FONTS = {
    'zh': dict(body='宋体', head='黑体', indent=21),   # 首行缩进两个汉字
    'en': dict(body='Times New Roman', head='Times New Roman', indent=18),
}
LANG = 'zh'


def set_cn_font(run, name=None, size=10.5, bold=False):
    if name is None:
        name = FONTS[LANG]['body']
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    # python-docx 不会自动设置东亚字体，需直接改底层 XML，否则中文会回退到默认字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
        name)


MIN_COL_PT = 30          # 数值列的最小宽度，避免被长文本列挤扁


def fit_columns(doc, table, header, rows):
    """按各列内容的最大显示宽度分配列宽，总宽严格等于正文宽度。

    python-docx 默认给每列等宽：表 3 的 "RF-DETR-Nano" 会在 61.7pt 的列里
    折成两行，而同表的数值列大量留白。改用 Word 的 autofit 则会把表撑到
    442.8pt，超出 432pt 的正文宽度——两端出血比折行更糟。
    故在此自行分配：按内容定比例，再缩放到正好填满正文宽度。
    """
    sec = doc.sections[-1]
    total = sec.page_width - sec.left_margin - sec.right_margin

    def disp_width(s):
        """显示宽度。汉字按两个西文字符计，行内标记不计入。"""
        s = re.sub(r'\*\*|`|</?sup>|\*', '', s)
        return sum(2 if is_cjk(c) else 1 for c in s)

    raw = []
    for k in range(len(header)):
        vals = [disp_width(header[k])]
        vals += [disp_width(r[k]) for r in rows if k < len(r)]
        raw.append(max(vals) + 2)        # +2 留作单元格内边距

    # 单列过长会吃掉其余列，封顶为均值的 3 倍
    cap = 3 * sum(raw) / len(raw)
    raw = [min(v, cap) for v in raw]

    scale = total / sum(raw)
    widths = [v * scale for v in raw]

    # 保底宽度：不足者补齐，缺口按比例从有余量的列扣回
    short = [i for i, v in enumerate(widths) if v < Pt(MIN_COL_PT)]
    if short:
        deficit = sum(Pt(MIN_COL_PT) - widths[i] for i in short)
        donors = [i for i in range(len(widths)) if i not in short]
        pool = sum(widths[i] for i in donors)
        for i in short:
            widths[i] = Pt(MIN_COL_PT)
        for i in donors:
            widths[i] -= deficit * (widths[i] / pool)

    table.autofit = False
    tblPr = table._tbl.tblPr
    for tag in ('tblLayout', 'tblW'):
        el = tblPr.find(qn('w:' + tag))
        if el is not None:
            tblPr.remove(el)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)

    # Word 只认单元格上的宽度，逐格写入
    for k, wd in enumerate(widths):
        for row in table.rows:
            row.cells[k].width = int(wd)


INLINE = re.compile(r'(\*\*[^*]+\*\*|<sup>.+?</sup>|`[^`]+`|\*[^*\n]+\*)')


def add_inline(par, text, size=10.5, base_bold=False):
    """处理 **加粗**、*斜体*、`代码` 与 <sup>上标</sup> 四种行内标记。

    上标用于参考文献的引用编号；斜体用于西文刊名。
    """
    for seg in INLINE.split(text):
        if not seg:
            continue
        if seg.startswith('**') and seg.endswith('**'):
            set_cn_font(par.add_run(seg[2:-2]), size=size, bold=True)
        elif seg.startswith('<sup>') and seg.endswith('</sup>'):
            r = par.add_run(seg[5:-6])
            set_cn_font(r, size=size, bold=base_bold)
            r.font.superscript = True
        elif seg.startswith('`') and seg.endswith('`'):
            r = par.add_run(seg[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(size - 0.5)
            r.font.color.rgb = RGBColor(0x88, 0x22, 0x22)
        elif seg.startswith('*') and seg.endswith('*'):
            r = par.add_run(seg[1:-1])
            set_cn_font(r, size=size, bold=base_bold)
            r.italic = True
        else:
            set_cn_font(par.add_run(seg), size=size, bold=base_bold)


def is_cjk(ch):
    return '一' <= ch <= '鿿' or '　' <= ch <= '〿'


def join_wrapped(lines):
    """还原被硬换行拆开的句子。

    中文两行之间直接相接；只要有一侧是西文（参考文献、英文术语），
    就补一个空格，否则会把 "networks." 与 "Nature" 粘成一个词。
    """
    out = ''
    for seg in lines:
        if not out:
            out = seg
            continue
        out += ('' if is_cjk(out[-1]) and is_cjk(seg[0]) else ' ') + seg
    return out


def is_table_sep(line):
    """markdown 表格的分隔行，如 |---|---|"""
    return bool(re.fullmatch(r'\|[\s\-:|]+\|', line.strip()))


def split_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def convert(md_paths, out_path):
    """md_paths 可以是一个路径，也可以是多个——多个时依次拼接并在之间分页。

    arXiv 要求多语言投稿合为单一文件且英文在前，故 English + Chinese 两份
    markdown 可直接一起传入。
    """
    if isinstance(md_paths, (str, Path)):
        md_paths = [md_paths]

    global LANG
    LANG = detect_lang(md_paths[0])
    doc = Document()
    doc.styles['Normal'].font.name = FONTS[LANG]['body']
    doc.styles['Normal'].font.size = Pt(10.5)

    for n, md_path in enumerate(md_paths):
        if n:
            doc.add_page_break()
        LANG = detect_lang(md_path)      # 合订本里两种语言各用各的字体
        _render(doc, md_path)

    doc.save(out_path)
    print(f'已生成 {out_path}')


def detect_lang(md_path):
    """按汉字占比判定语言。合订本中每份文件各自判定。"""
    t = open(md_path, encoding='utf-8').read()
    cjk = sum('一' <= c <= '鿿' for c in t)
    return 'zh' if cjk / max(len(t), 1) > 0.05 else 'en'


def _render(doc, md_path):
    lines = open(md_path, encoding='utf-8').read().split('\n')

    i = 0
    while i < len(lines):
        ln = lines[i]
        st = ln.strip()

        # ---- 表格：当前行以 | 开头且下一行是分隔行 ----
        if st.startswith('|') and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            header = split_row(st)
            rows = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith('|'):
                rows.append(split_row(lines[j]))
                j += 1
            t = doc.add_table(rows=1, cols=len(header))
            t.style = 'Table Grid'
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for k, h in enumerate(header):
                cell = t.rows[0].cells[k]
                cell.text = ''
                add_inline(cell.paragraphs[0], h, size=9.5, base_bold=True)
            for r in rows:
                cells = t.add_row().cells
                for k, v in enumerate(r[:len(header)]):
                    cells[k].text = ''
                    add_inline(cells[k].paragraphs[0], v, size=9.5)
            fit_columns(doc, t, header, rows)
            doc.add_paragraph()
            i = j
            continue

        # ---- 图片：![alt](相对路径) 独占一行 ----
        m = re.fullmatch(r'!\[[^\]]*\]\(([^)]+)\)', st)
        if m:
            img = (Path(md_path).parent / m.group(1)).resolve()
            if img.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 正文宽度约 15.9 cm（A4 减默认页边距），留出余量
                p.add_run().add_picture(str(img), width=Cm(13.5))
            else:
                print(f'  [警告] 图片不存在，已跳过: {img}')
            i += 1
            continue

        # ---- 标题 ----
        m = re.match(r'^(#{1,4})\s+(.*)', st)
        if m:
            lvl, txt = len(m.group(1)), m.group(2)
            p = doc.add_paragraph()
            sizes = {1: 16, 2: 14, 3: 12, 4: 11}
            set_cn_font(p.add_run(txt), name=FONTS[LANG]['head'],
                        size=sizes.get(lvl, 11), bold=True)
            if lvl == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # ---- 分隔线 ----
        if st in ('---', '***', '___'):
            i += 1
            continue

        # ---- 引用块：合并连续的 > 行 ----
        if st.startswith('>'):
            buf = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                buf.append(lines[i].strip().lstrip('>').strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_inline(p, ' '.join(x for x in buf if x), size=9.5)
            continue

        # ---- 无序列表 ----
        if re.match(r'^[-*]\s+', st):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, re.sub(r'^[-*]\s+', '', st))
            i += 1
            continue

        # ---- 有序列表 ----
        if re.match(r'^\d+\.\s+', st):
            p = doc.add_paragraph(style='List Number')
            add_inline(p, re.sub(r'^\d+\.\s+', '', st))
            i += 1
            continue

        # ---- 空行 ----
        if not st:
            i += 1
            continue

        # ---- 正文段落：合并到下一个空行为止，还原被硬换行拆开的句子 ----
        buf = [st]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(('#', '|', '>', '---'))
                    or re.match(r'^[-*]\s+|^\d+\.\s+', nxt)):
                break
            buf.append(nxt)
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(FONTS[LANG]['indent'])
        add_inline(p, join_wrapped(buf))


if __name__ == '__main__':
    ROOT = r'D:\dev\crop-detect'
    if len(sys.argv) > 2:
        convert(sys.argv[1:-1], sys.argv[-1])
    else:
        convert(rf'{ROOT}\paper\论文.md', rf'{ROOT}\paper\论文.docx')
